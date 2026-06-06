#!/usr/bin/env python3
"""
OSINT Deep Briefing Engine — Content Writer Edition
=====================================================
Auto-searches the internet, analyses via AI, outputs investigative briefings.
API keys are loaded from .env — NEVER hardcode keys in this file.

Usage:
  python osint_bot.py --target "Entity Name"
  python osint_bot.py --target "Entity" --provider groq --mode article
  python osint_bot.py --target "Entity" --angle funding --quotes --seo
  python osint_bot.py --target "Entity" --legal-review --export docx
  python osint_bot.py --target "Entity" --watchlist add
  python osint_bot.py --target "Entity" --watchlist run
  python osint_bot.py --providers
"""

import os
import sys
import re
import json
import time
import argparse
import textwrap
import hashlib
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
from rich.table import Table
from rich.rule import Rule
from rich import print as rprint

# Load .env file — keys stay out of code
load_dotenv()

console = Console()

# ─────────────────────────────────────────────────────────────────────────────
# PROVIDER REGISTRY
# ─────────────────────────────────────────────────────────────────────────────

PROVIDERS = {
    "groq": {
        "name": "Groq (FREE ✓)",
        "env_key": "GROQ_API_KEY",
        "model": "llama-3.3-70b-versatile",
        "free": True,
        "get_key_url": "https://console.groq.com → API Keys  (no credit card needed)",
    },
    "gemini": {
        "name": "Google Gemini (FREE ✓)",
        "env_key": "GEMINI_API_KEY",
        "model": "gemini-2.0-flash",
        "free": True,
        "get_key_url": "https://aistudio.google.com/app/apikey  (no credit card needed)",
    },
    "claude": {
        "name": "Anthropic Claude",
        "env_key": "ANTHROPIC_API_KEY",
        "model": "claude-sonnet-4-20250514",
        "free": False,
        "get_key_url": "https://console.anthropic.com → API Keys",
    },
    "openai": {
        "name": "OpenAI GPT-4o",
        "env_key": "OPENAI_API_KEY",
        "model": "gpt-4o",
        "free": False,
        "get_key_url": "https://platform.openai.com/api-keys",
    },
}


class UsageFake:
    """Dummy usage object for providers that don't return token counts."""
    def __init__(self):
        self.input_tokens = "n/a"
        self.output_tokens = "n/a"


def call_llm(system: str, user: str, provider: str, api_key: str,
             max_tokens: int = 6000) -> tuple[str, object]:
    """Universal LLM caller — routes to the correct provider SDK."""

    if provider == "claude":
        import anthropic as _ant
        client = _ant.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model=PROVIDERS["claude"]["model"],
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": user}],
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
        )
        text = "".join(b.text for b in msg.content if hasattr(b, "text"))
        return text, msg.usage

    elif provider == "groq":
        from openai import OpenAI
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.groq.com/openai/v1",
        )
        resp = client.chat.completions.create(
            model=PROVIDERS["groq"]["model"],
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        )
        return resp.choices[0].message.content, UsageFake()

    elif provider == "gemini":
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            model_name=PROVIDERS["gemini"]["model"],
            system_instruction=system,
        )
        resp = model.generate_content(
            user,
            generation_config={"max_output_tokens": max_tokens, "temperature": 0.3},
        )
        return resp.text, UsageFake()

    elif provider == "openai":
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=PROVIDERS["openai"]["model"],
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        )
        return resp.choices[0].message.content, UsageFake()

    else:
        raise ValueError(f"Unknown provider: {provider}")


# ─────────────────────────────────────────────────────────────────────────────
# SYSTEM PROMPTS
# ─────────────────────────────────────────────────────────────────────────────

BRIEFING_SYSTEM = """You are an elite Investigative Journalist and OSINT Expert with internet search capability.

Your task: Research the target entity thoroughly using web search (if available) or your knowledge, and produce a structured intelligence briefing.

SEARCH STRATEGY:
1. Search entity official presence, background, founding, key people
2. Search news, investigative reports, controversies
3. Search financial records, funding, ownership structure
4. Search legal cases, regulatory filings, court records
5. Search social media, public statements, deleted content
6. Search contradictions between official narrative and public record

Generate a Deep Intelligence Briefing as a valid JSON object with exactly these keys:
- "context": array of bullet strings — surface claims vs hidden agenda, launch timeline, triggers
- "network": array of bullet strings — founders, directors, promoters, their political/corporate/institutional ties
- "trail": array of bullet strings — funding, parent companies, domain ownership, operational infrastructure
- "guns": array of bullet strings — contradictions, narrative shifts, hypocrisy, deleted post leaks
- "quotes": array of objects {speaker, quote, date, source_url, is_direct}
- "confidence": object mapping "context"|"network"|"trail"|"guns" to "high"|"medium"|"low"
- "needs_corroboration": array of strings
- "named_persons": array of strings — full names of all named individuals
- "keywords": array of 8-12 SEO keyword strings
- "headline_variants": object with keys "seo", "clickable", "print", "thread_hook"
- "sources": array of {index, title, url, domain, type, snippet} — all sources used

CRITICAL RULES:
1. STRICT FACTUAL ACCURACY — only facts from web search or verified knowledge. No hallucination.
2. EXPLICIT CITATIONS — every bullet MUST end with [Source N] matching sources array index.
3. UNKNOWN HANDLING — write "NO DIRECT DATA FOUND IN CURRENT SOURCES" if absent.
4. confidence: high = multiple corroborating sources, medium = single source, low = inferred.
5. is_direct=true only if actual quotation marks are used in the source.
6. Include 6-15 diverse sources with real URLs.

Respond ONLY with a valid JSON object. No markdown fences, no preamble."""


ARTICLE_SYSTEM = """You are a senior investigative journalist writing for a national publication.
Given a structured OSINT briefing JSON and target entity name, write a complete investigative article.

Structure:
- HEADLINE: compelling, factual headline
- SUBHEADLINE: one-sentence summary
- BYLINE: "By [Investigative Desk]"
- DATELINE: today's date
- LEDE (2-3 sentences): the most newsworthy revelation
- BODY (5-8 paragraphs): weave context, network, money trail, contradictions into narrative prose
- SIDEBAR: "Key Facts at a Glance" — 5 bullet points
- CLOSING: future implications or unanswered questions

Rules:
- Use "alleged", "according to", "reportedly" for unverified claims
- Attribute every specific claim to its source
- Never editorialize without attribution
- Return as plain text (no markdown)"""


THREAD_SYSTEM = """You are a viral investigative journalist on X (Twitter).
Given an OSINT briefing JSON, write a numbered thread that exposes the story.

Rules:
- Tweet 1: the bombshell hook — max 280 chars
- Tweets 2-10: one revelation per tweet, build tension progressively
- Each tweet max 280 characters
- End with a THREAD SUMMARY tweet and a CTA
- No hashtags unless requested
- Return as plain numbered list"""


NEWSLETTER_SYSTEM = """You are writing for an investigative newsletter with 50,000 subscribers.
Given an OSINT briefing JSON, write a newsletter edition.

Structure:
- SUBJECT LINE: (email subject, max 60 chars)
- PREVIEW TEXT: (email preview, max 90 chars)
- GREETING: "Dear Reader,"
- INTRO (2 sentences): why this story matters today
- THE STORY (4-6 paragraphs): narrative, readable, gripping
- WHAT TO WATCH: 3 bullet points of follow-up angles
- SIGN-OFF: "Until next time, [The Investigative Desk]"

Tone: smart, trustworthy, slightly urgent. No jargon. Return as plain text."""


LEGAL_SYSTEM = """You are a media lawyer and editorial compliance expert.
Given investigative content, rewrite it with legal risk mitigation applied.

Rules:
- Replace definitive allegations about named individuals with "alleged", "according to [source]", "reportedly"
- Flag statements that could constitute defamation with [LEGAL-FLAG: reason]
- Add "did not respond to requests for comment" where relevant
- Preserve all well-sourced factual claims unchanged
- Return as plain text with legal flags visible"""


# ─────────────────────────────────────────────────────────────────────────────
# SCRAPER (optional — used when user provides explicit URLs)
# ─────────────────────────────────────────────────────────────────────────────

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "en-US,en;q=0.9",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
}

WAYBACK_API = "http://archive.org/wayback/available?url={url}"


def fetch_wayback(url: str) -> str | None:
    try:
        r = requests.get(WAYBACK_API.format(url=url), timeout=10)
        data = r.json()
        snapshot = data.get("archived_snapshots", {}).get("closest", {})
        if snapshot.get("available"):
            return snapshot["url"]
    except Exception:
        pass
    return None


def fetch_youtube_transcript(url: str) -> str | None:
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        vid_match = re.search(r"(?:v=|youtu\.be/)([A-Za-z0-9_-]{11})", url)
        if not vid_match:
            return None
        vid_id = vid_match.group(1)
        transcript = YouTubeTranscriptApi.get_transcript(vid_id)
        return " ".join(t["text"] for t in transcript)[:15000]
    except Exception:
        return None


def scrape_url(url: str, index: int, timeout: int = 15) -> dict:
    result = {
        "index": index, "url": url,
        "domain": urlparse(url).netloc,
        "title": "", "text": "", "status": "ok",
        "error": None, "char_count": 0, "source_type": "web",
        "wayback_used": False,
    }

    if "youtube.com" in url or "youtu.be" in url:
        transcript = fetch_youtube_transcript(url)
        if transcript:
            result["text"] = transcript
            result["char_count"] = len(transcript)
            result["title"] = f"YouTube Video: {url}"
            result["source_type"] = "youtube_transcript"
            return result
        result["status"] = "error"
        result["error"] = "YouTube transcript unavailable"
        return result

    def _fetch(target_url: str) -> requests.Response:
        return requests.get(target_url, headers=HEADERS, timeout=timeout)

    try:
        resp = _fetch(url)
        resp.raise_for_status()
    except (requests.exceptions.HTTPError, requests.exceptions.ConnectionError,
            requests.exceptions.Timeout) as e:
        wb_url = fetch_wayback(url)
        if wb_url:
            try:
                resp = _fetch(wb_url)
                resp.raise_for_status()
                result["wayback_used"] = True
                result["source_type"] = "wayback"
            except Exception as e2:
                result["status"] = "error"
                result["error"] = f"{e} | Wayback also failed: {e2}"
                return result
        else:
            result["status"] = "error"
            result["error"] = str(e)
            return result
    except Exception as e:
        result["status"] = "error"
        result["error"] = str(e)
        return result

    try:
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header",
                         "aside", "form", "noscript", "iframe"]):
            tag.decompose()

        result["title"] = soup.title.string.strip() if soup.title else url
        main = soup.find("article") or soup.find("main") or soup.find("body")
        if main:
            paras = main.find_all(["p", "h1", "h2", "h3", "h4", "li",
                                   "blockquote", "td", "th"])
            lines = [p.get_text(separator=" ", strip=True)
                     for p in paras if len(p.get_text(strip=True)) > 30]
            result["text"] = "\n".join(lines)[:15000]
        else:
            result["text"] = soup.get_text(separator="\n", strip=True)[:15000]

        result["char_count"] = len(result["text"])
    except Exception as e:
        result["status"] = "parse_error"
        result["error"] = str(e)

    return result


def scrape_all(urls: list[str], timeout: int = 15) -> list[dict]:
    results = []
    with Progress(
        SpinnerColumn(),
        TextColumn("[bold red]Scraping[/bold red] {task.description}"),
        BarColumn(),
        TextColumn("[dim]{task.completed}/{task.total}[/dim]"),
        console=console,
    ) as progress:
        task = progress.add_task("sources...", total=len(urls))
        for i, url in enumerate(urls):
            progress.update(task, description=f"[dim]{urlparse(url).netloc}[/dim]")
            results.append(scrape_url(url, i, timeout=timeout))
            time.sleep(0.6)
            progress.advance(task)
    return results


# ─────────────────────────────────────────────────────────────────────────────
# PROMPT BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def build_briefing_prompt_no_urls(target: str, extra: str, angle: str | None,
                                   depth: str) -> str:
    """Prompt for internet-search mode — no manual URLs needed."""
    depth_instr = {
        "standard": "Standard investigative briefing.",
        "deep": "Deep dive — include historical records, archived pages, secondary sources.",
        "financial": "Focus on financial records, funding, investors, revenue, controversies.",
        "network": "Focus on people network — founders, directors, advisors, political ties.",
    }.get(depth, "Standard investigative briefing.")

    lines = [
        f"TARGET ENTITY: {target}",
        "",
        f"DEPTH: {depth_instr}",
        "",
    ]
    if angle and angle != "all":
        lines += [f"FOCUS ANGLE: Prioritise findings related to '{angle}'.", ""]
    if extra.strip():
        lines += ["EXTRA CONTEXT / MANUAL NOTES:", extra.strip(), ""]
    lines.append(
        "Search the internet comprehensively for this entity and produce "
        "the full Deep Intelligence Briefing JSON now."
    )
    return "\n".join(lines)


def build_briefing_prompt_with_urls(target: str, scraped: list[dict], extra: str,
                                     angle: str | None, depth: str) -> str:
    """Prompt when user has provided explicit URLs to scrape."""
    depth_instr = {
        "standard": "Standard investigative briefing.",
        "deep": "Deep dive — include all details from provided sources.",
        "financial": "Focus on financial findings from the provided sources.",
        "network": "Focus on people/network findings from the provided sources.",
    }.get(depth, "Standard investigative briefing.")

    lines = [
        f"TARGET ENTITY: {target}",
        "",
        f"DEPTH: {depth_instr}",
        "",
    ]
    if angle and angle != "all":
        lines += [f"FOCUS ANGLE: Prioritise findings related to '{angle}'.", ""]

    lines += ["RAW RESEARCH DATA FROM PROVIDED URLS:", ""]
    for s in scraped:
        lines.append(f"[Source {s['index']}] URL: {s['url']}")
        lines.append(f"Domain: {s['domain']}")
        if s.get("wayback_used"):
            lines.append("Note: fetched from Wayback Machine archive")
        if s["title"]:
            lines.append(f"Title: {s['title']}")
        if s["status"] == "ok":
            lines.append(f"Content ({s['char_count']} chars):\n{s['text']}")
        else:
            lines.append(f"[SCRAPE FAILED: {s['error']}]")
        lines += ["", "─" * 60, ""]

    if extra.strip():
        lines += ["EXTRA CONTEXT / MANUAL NOTES:", extra.strip(), ""]

    lines.append("Generate the full Deep Intelligence Briefing JSON now.")
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# LLM CALL WRAPPERS
# ─────────────────────────────────────────────────────────────────────────────

def parse_json_response(raw: str) -> dict:
    clean = re.sub(r"```json|```", "", raw).strip()
    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", clean)
        if match:
            return json.loads(match.group())
        raise ValueError(f"Could not parse JSON.\nRaw (first 500):\n{clean[:500]}")


def run_briefing(target: str, prompt: str, provider: str,
                 api_key: str) -> tuple[dict, object]:
    pname = PROVIDERS[provider]["name"]
    console.print(f"\n[bold red]◌[/bold red] Analysing with {pname}...", end=" ")
    raw, usage = call_llm(BRIEFING_SYSTEM, prompt, provider, api_key)
    console.print("[green]DONE[/green]")
    return parse_json_response(raw), usage


def run_article(briefing: dict, target: str, tone: str,
                provider: str, api_key: str) -> str:
    console.print("[bold red]◌[/bold red] Drafting article...", end=" ")
    user = (f"TARGET: {target}\nTONE: {tone}\n\n"
            f"BRIEFING JSON:\n{json.dumps(briefing, indent=2)}")
    raw, _ = call_llm(ARTICLE_SYSTEM, user, provider, api_key, max_tokens=3000)
    console.print("[green]DONE[/green]")
    return raw


def run_thread(briefing: dict, target: str, provider: str, api_key: str) -> str:
    console.print("[bold red]◌[/bold red] Writing thread...", end=" ")
    user = f"TARGET: {target}\n\nBRIEFING JSON:\n{json.dumps(briefing, indent=2)}"
    raw, _ = call_llm(THREAD_SYSTEM, user, provider, api_key, max_tokens=2000)
    console.print("[green]DONE[/green]")
    return raw


def run_newsletter(briefing: dict, target: str, provider: str, api_key: str) -> str:
    console.print("[bold red]◌[/bold red] Writing newsletter...", end=" ")
    user = f"TARGET: {target}\n\nBRIEFING JSON:\n{json.dumps(briefing, indent=2)}"
    raw, _ = call_llm(NEWSLETTER_SYSTEM, user, provider, api_key, max_tokens=2500)
    console.print("[green]DONE[/green]")
    return raw


def run_legal_review(content: str, provider: str, api_key: str) -> str:
    console.print("[bold red]◌[/bold red] Legal compliance review...", end=" ")
    raw, _ = call_llm(LEGAL_SYSTEM, content, provider, api_key, max_tokens=3000)
    console.print("[green]DONE[/green]")
    return raw


# ─────────────────────────────────────────────────────────────────────────────
# WATCHLIST
# ─────────────────────────────────────────────────────────────────────────────

WATCHLIST_FILE = Path("./output/watchlist.json")


def watchlist_load() -> dict:
    if WATCHLIST_FILE.exists():
        return json.loads(WATCHLIST_FILE.read_text())
    return {}


def watchlist_save(data: dict):
    WATCHLIST_FILE.parent.mkdir(parents=True, exist_ok=True)
    WATCHLIST_FILE.write_text(json.dumps(data, indent=2))


def watchlist_add(target: str, urls: list[str]):
    wl = watchlist_load()
    wl[target] = {
        "urls": urls,
        "added": datetime.now().isoformat(),
        "runs": [],
    }
    watchlist_save(wl)
    console.print(
        f"[green]✓[/green] Added [bold]{target}[/bold] to watchlist "
        f"with {len(urls)} URL(s)"
    )


def watchlist_list():
    wl = watchlist_load()
    if not wl:
        console.print("[dim]Watchlist is empty.[/dim]")
        return
    t = Table(show_header=True, header_style="bold dim", box=None)
    t.add_column("Entity", style="white", width=30)
    t.add_column("URLs", width=8)
    t.add_column("Runs", width=8)
    t.add_column("Last run", width=20)
    for name, data in wl.items():
        last = data["runs"][-1]["at"] if data["runs"] else "Never"
        t.add_row(name, str(len(data["urls"])), str(len(data["runs"])), last)
    console.print(t)


def compute_delta(old_briefing: dict, new_briefing: dict) -> dict:
    delta = {}
    for key in ["context", "network", "trail", "guns"]:
        old_set = set(old_briefing.get(key, []))
        new_set = set(new_briefing.get(key, []))
        added = list(new_set - old_set)
        removed = list(old_set - new_set)
        if added or removed:
            delta[key] = {"added": added, "removed": removed}
    return delta


# ─────────────────────────────────────────────────────────────────────────────
# TERMINAL RENDERER
# ─────────────────────────────────────────────────────────────────────────────

SECTION_META = [
    ("01", "context", "THE REIGNING CONTEXT & COVERT IDENTITY",      "red"),
    ("02", "network", "THE HUMAN NETWORK & AFFILIATIONS (KUNDALI)",   "yellow"),
    ("03", "trail",   "THE PAPER TRAIL & MONEY TRAIL",                "cyan"),
    ("04", "guns",    "SMOKING GUNS & CONTRADICTIONS",                "magenta"),
]

CONFIDENCE_ICON = {
    "high":   "[green]●[/green]",
    "medium": "[yellow]●[/yellow]",
    "low":    "[red]●[/red]",
}


def render_terminal(target: str, briefing: dict, sources: list[dict],
                    usage, show_quotes: bool = False, show_seo: bool = False):
    console.print()
    console.rule("[bold red]DEEP INTELLIGENCE BRIEFING[/bold red]", style="red")
    confidence = briefing.get("confidence", {})
    console.print(Panel(
        f"[bold white]ENTITY:[/bold white] {target}\n"
        f"[dim]Generated: {datetime.now().strftime('%d %b %Y · %H:%M:%S')}  |  "
        f"Tokens: {usage.input_tokens} in / {usage.output_tokens} out[/dim]",
        border_style="red", expand=True,
    ))
    console.print()

    for num, key, title, color in SECTION_META:
        conf = confidence.get(key, "medium")
        conf_badge = CONFIDENCE_ICON.get(conf, "")
        console.print(
            f"[bold {color}]### {num}. {title}[/bold {color}]  "
            f"{conf_badge} [dim]{conf.upper()} CONFIDENCE[/dim]"
        )
        console.print()
        for bullet in briefing.get(key, ["No data found."]):
            fmt = re.sub(
                r"\[Source (\d+)\]",
                r"[dim cyan]\[Source \1\][/dim cyan]",
                bullet,
            )
            fmt = fmt.replace(
                "NO DIRECT DATA FOUND IN CURRENT SOURCES",
                "[bold yellow]⚑ NO DIRECT DATA FOUND IN CURRENT SOURCES[/bold yellow]",
            )
            wrapped = textwrap.fill(fmt, width=110, subsequent_indent="  ")
            console.print(f"  [bold {color}]▸[/bold {color}] {wrapped}")
            console.print()
        console.rule(style="dim")
        console.print()

    needs = briefing.get("needs_corroboration", [])
    if needs:
        console.print("[bold yellow]⚠  NEEDS CORROBORATION BEFORE PUBLISHING[/bold yellow]")
        for item in needs:
            console.print(f"  [yellow]▸[/yellow] {item}")
        console.print()

    persons = briefing.get("named_persons", [])
    if persons:
        console.print(
            "[bold]RIGHT-OF-REPLY CHECKLIST[/bold]  "
            "[dim](contact before publishing)[/dim]"
        )
        for p in persons:
            console.print(f"  [dim]☐[/dim]  {p}")
        console.print()

    if show_quotes:
        quotes = briefing.get("quotes", [])
        if quotes:
            console.print("[bold cyan]QUOTES BANK[/bold cyan]")
            for q in quotes:
                direct_tag = (
                    "[green](DIRECT)[/green]"
                    if q.get("is_direct")
                    else "[yellow](PARAPHRASE)[/yellow]"
                )
                console.print(
                    f"  {direct_tag} [bold]{q.get('speaker', 'Unknown')}[/bold]"
                    f" — {q.get('date', 'n/d')}"
                )
                console.print(f'  [italic]"{q.get("quote", "")}"[/italic]')
                console.print(f"  [dim]Source: {q.get('source_url', '')}[/dim]")
                console.print()

    if show_seo:
        headlines = briefing.get("headline_variants", {})
        keywords = briefing.get("keywords", [])
        console.print("[bold green]SEO & HEADLINES[/bold green]")
        for variant, label in [
            ("seo", "SEO"), ("clickable", "Clickable"),
            ("print", "Print"), ("thread_hook", "Thread Hook"),
        ]:
            if headlines.get(variant):
                console.print(f"  [dim]{label}:[/dim] {headlines[variant]}")
        if keywords:
            console.print(f"  [dim]Keywords:[/dim] {', '.join(keywords)}")
        console.print()

    # Sources table
    src_list = sources or briefing.get("sources", [])
    if src_list:
        console.print("[bold]SOURCES[/bold]")
        table = Table(show_header=True, header_style="bold dim", box=None, padding=(0, 2))
        table.add_column("ID", style="cyan", width=6)
        table.add_column("Type", width=12)
        table.add_column("Domain", width=28)
        table.add_column("Status", width=10)
        table.add_column("URL", style="dim", max_width=60)
        for s in src_list:
            # scraped sources have status; briefing sources don't
            status_str = (
                "[green]OK[/green]"
                if s.get("status", "ok") == "ok"
                else f"[red]{s.get('status', 'ok')}[/red]"
            )
            wb = " [dim](wb)[/dim]" if s.get("wayback_used") else ""
            table.add_row(
                f"[{s.get('index', '?')}]",
                s.get("source_type", s.get("type", "web")),
                s.get("domain", urlparse(s.get("url", "")).netloc) + wb,
                status_str,
                s.get("url", ""),
            )
        console.print(table)
        console.print()


def render_delta(target: str, delta: dict):
    console.print()
    console.rule(f"[bold yellow]DELTA REPORT — {target}[/bold yellow]", style="yellow")
    if not delta:
        console.print("[green]No changes detected since last run.[/green]")
        return
    for key, changes in delta.items():
        title = next(t for _, k, t, _ in SECTION_META if k == key)
        if changes["added"]:
            console.print(f"[bold green]NEW in {title}:[/bold green]")
            for item in changes["added"]:
                console.print(f"  [green]+[/green] {item}")
        if changes["removed"]:
            console.print(f"[bold red]REMOVED from {title}:[/bold red]")
            for item in changes["removed"]:
                console.print(f"  [red]−[/red] {item}")
        console.print()


# ─────────────────────────────────────────────────────────────────────────────
# FILE SAVERS
# ─────────────────────────────────────────────────────────────────────────────

def _make_filename(target: str, suffix: str, ext: str, output_dir: Path) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = re.sub(r"[^\w\s-]", "", target).strip().replace(" ", "_")[:40]
    return output_dir / f"{safe}_{suffix}_{ts}.{ext}"


def save_briefing_markdown(target: str, briefing: dict, sources: list[dict],
                            output_dir: Path, include_quotes: bool = False,
                            include_seo: bool = False) -> Path:
    filepath = _make_filename(target, "briefing", "md", output_dir)
    lines = [
        "# DEEP INTELLIGENCE BRIEFING", "",
        f"**Entity:** {target}  ",
        f"**Generated:** {datetime.now().strftime('%d %b %Y at %H:%M:%S')}  ",
        "**Classification:** CONFIDENTIAL — FOR RESEARCH USE ONLY",
        "", "---", "",
    ]
    confidence = briefing.get("confidence", {})
    for num, key, title, _ in SECTION_META:
        conf = confidence.get(key, "medium").upper()
        lines += [f"## {num}. {title}  *(Confidence: {conf})*", ""]
        for b in briefing.get(key, ["No data found."]):
            lines.append(f"- {b}")
        lines += ["", "---", ""]

    needs = briefing.get("needs_corroboration", [])
    if needs:
        lines += ["## ⚠ NEEDS CORROBORATION", ""]
        for n in needs:
            lines.append(f"- {n}")
        lines += ["", "---", ""]

    persons = briefing.get("named_persons", [])
    if persons:
        lines += ["## RIGHT-OF-REPLY CHECKLIST", ""]
        for p in persons:
            lines.append(f"- [ ] {p}")
        lines += ["", "---", ""]

    if include_quotes:
        quotes = briefing.get("quotes", [])
        if quotes:
            lines += ["## QUOTES BANK", ""]
            for q in quotes:
                tag = "(DIRECT)" if q.get("is_direct") else "(PARAPHRASE)"
                lines.append(
                    f"### {q.get('speaker', 'Unknown')} — "
                    f"{q.get('date', 'n/d')} {tag}"
                )
                lines.append(f"> {q.get('quote', '')}")
                lines.append(f"*Source: {q.get('source_url', '')}*")
                lines.append("")
            lines += ["---", ""]

    if include_seo:
        headlines = briefing.get("headline_variants", {})
        keywords = briefing.get("keywords", [])
        lines += ["## SEO & HEADLINES", ""]
        for k, v in headlines.items():
            lines.append(f"**{k.upper()}:** {v}")
        lines.append(f"\n**Keywords:** {', '.join(keywords)}")
        lines += ["", "---", ""]

    src_list = sources or briefing.get("sources", [])
    lines += ["## SOURCES", ""]
    for s in src_list:
        domain = s.get("domain", urlparse(s.get("url", "")).netloc)
        wb = " *(Wayback)*" if s.get("wayback_used") else ""
        ok = "✓" if s.get("status", "ok") == "ok" else "✗"
        lines.append(
            f"- **[Source {s.get('index', '?')}]** {ok}{wb} "
            f"[{s.get('title', domain)}]({s.get('url', '')})"
        )
        if s.get("error"):
            lines.append(f"  - Error: {s['error']}")

    lines += [
        "", "---",
        "*All facts sourced strictly from provided URLs or verified knowledge. "
        "Verify independently before publication.*",
    ]
    filepath.write_text("\n".join(lines), encoding="utf-8")
    return filepath


def save_text(target: str, content: str, suffix: str, output_dir: Path) -> Path:
    filepath = _make_filename(target, suffix, "txt", output_dir)
    filepath.write_text(content, encoding="utf-8")
    return filepath


def save_json_output(target: str, briefing: dict, sources: list[dict],
                     output_dir: Path, redact_sources: bool = False) -> Path:
    filepath = _make_filename(target, "data", "json", output_dir)
    src_out = []
    for s in sources:
        entry = dict(s)
        if redact_sources:
            entry["url"] = (
                hashlib.sha256(s["url"].encode()).hexdigest()[:16] + "...[REDACTED]"
            )
            entry["text"] = "[REDACTED]"
        src_out.append(entry)

    payload = {
        "target": target,
        "generated_at": datetime.now().isoformat(),
        "briefing": briefing,
        "sources": src_out or briefing.get("sources", []),
    }
    filepath.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return filepath


def save_docx(target: str, briefing: dict, sources: list[dict],
              output_dir: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("DEEP INTELLIGENCE BRIEFING", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("Entity: ").bold = True
    meta.add_run(target)
    meta.add_run(f"\nGenerated: {datetime.now().strftime('%d %B %Y at %H:%M')}")
    run = meta.add_run("\nClassification: CONFIDENTIAL — FOR RESEARCH USE ONLY")
    run.bold = True

    doc.add_paragraph()
    confidence = briefing.get("confidence", {})

    for num, key, title_str, _ in SECTION_META:
        conf = confidence.get(key, "medium").upper()
        doc.add_heading(f"{num}. {title_str}  [Confidence: {conf}]", level=1)
        for b in briefing.get(key, ["No data found."]):
            doc.add_paragraph(b, style="List Bullet")

    needs = briefing.get("needs_corroboration", [])
    if needs:
        doc.add_heading("⚠ NEEDS CORROBORATION", level=1)
        for n in needs:
            doc.add_paragraph(n, style="List Bullet")

    persons = briefing.get("named_persons", [])
    if persons:
        doc.add_heading("RIGHT-OF-REPLY CHECKLIST", level=1)
        for p in persons:
            doc.add_paragraph(f"☐  {p}", style="List Bullet")

    src_list = sources or briefing.get("sources", [])
    doc.add_heading("SOURCES", level=1)
    for s in src_list:
        ok = "✓" if s.get("status", "ok") == "ok" else "✗"
        domain = s.get("domain", urlparse(s.get("url", "")).netloc)
        doc.add_paragraph(
            f"{ok} [Source {s.get('index', '?')}]  {domain}  {s.get('url', '')}"
        )

    filepath = _make_filename(target, "briefing", "docx", output_dir)
    doc.save(str(filepath))
    return filepath


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def parse_args():
    parser = argparse.ArgumentParser(
        description="OSINT Deep Briefing Engine — Content Writer Edition",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""
        FREE providers (no credit card needed):
          groq    → console.groq.com        (Llama 3.3 70B, very fast)
          gemini  → aistudio.google.com     (Gemini 2.0 Flash)

        Paid providers:
          claude  → console.anthropic.com   (Claude Sonnet — best quality, has live web search)
          openai  → platform.openai.com     (GPT-4o)

        API keys go in .env file — never hardcode them:
          GROQ_API_KEY=gsk_...
          GEMINI_API_KEY=AIza...
          ANTHROPIC_API_KEY=sk-ant-...
          OPENAI_API_KEY=sk-...

        Examples:
          python osint_bot.py --target "Byju's"
          python osint_bot.py --target "Adani Group" --provider gemini --mode article
          python osint_bot.py --target "ABC NGO" --provider groq --angle funding --quotes --seo
          python osint_bot.py --target "XYZ Media" --urls https://example.com --export docx
          python osint_bot.py --target "Entity" --watchlist add --urls url1 url2
          python osint_bot.py --target "Entity" --watchlist run
          python osint_bot.py --providers
        """),
    )

    parser.add_argument(
        "--provider",
        choices=list(PROVIDERS.keys()),
        default="groq",
        help="AI provider (default: groq — FREE, no credit card)",
    )
    parser.add_argument(
        "--providers",
        action="store_true",
        help="List all available providers and quit",
    )
    parser.add_argument("--target", help="Name of the entity being investigated")
    parser.add_argument("--urls", nargs="+", metavar="URL",
                        help="Optional: provide specific URLs to scrape")
    parser.add_argument("--file", metavar="FILE",
                        help="Optional: text file with one URL per line")
    parser.add_argument("--extra", default="", metavar="TEXT",
                        help="Extra context or manual notes to include")
    parser.add_argument(
        "--mode",
        choices=["briefing", "article", "thread", "newsletter"],
        default="briefing",
        help="Output mode (default: briefing)",
    )
    parser.add_argument(
        "--tone",
        choices=["formal", "tabloid", "academic"],
        default="formal",
    )
    parser.add_argument(
        "--depth",
        choices=["standard", "deep", "financial", "network"],
        default="standard",
        help="Investigation depth (default: standard)",
    )
    parser.add_argument(
        "--angle",
        choices=["all", "funding", "network", "contradictions", "legal", "timeline"],
        default="all",
        help="Focus angle for the investigation",
    )
    parser.add_argument("--quotes", action="store_true",
                        help="Extract and display quotes bank")
    parser.add_argument("--seo", action="store_true",
                        help="Generate SEO headlines and keywords")
    parser.add_argument("--legal-review", action="store_true", dest="legal_review",
                        help="Run legal compliance pass on the output")
    parser.add_argument(
        "--export",
        choices=["markdown", "docx", "json", "all"],
        default="markdown",
        help="Export format (default: markdown)",
    )
    parser.add_argument("--redact-sources", action="store_true", dest="redact_sources",
                        help="Hash/redact source URLs in JSON output")
    parser.add_argument(
        "--watchlist",
        choices=["add", "run", "list"],
        help="Watchlist: add entity, re-run entity, or list all",
    )
    parser.add_argument("--output-dir", default="./output", metavar="DIR")
    parser.add_argument("--api-key", metavar="KEY",
                        help="Override API key from .env (not recommended)")
    parser.add_argument("--timeout", type=int, default=15)
    return parser.parse_args()


def show_providers():
    console.print()
    t = Table(show_header=True, header_style="bold", box=None, padding=(0, 2))
    t.add_column("Flag", style="cyan", width=10)
    t.add_column("Provider", width=22)
    t.add_column("Model", width=28)
    t.add_column("Free?", width=8)
    t.add_column("Get your key at", style="dim")
    for flag, p in PROVIDERS.items():
        free = "[green]YES ✓[/green]" if p["free"] else "[yellow]Paid[/yellow]"
        t.add_row(flag, p["name"], p["model"], free, p["get_key_url"])
    console.print(t)
    console.print()
    console.print("[dim]Keys go in .env file:[/dim]")
    for _, p in PROVIDERS.items():
        console.print(f"  [dim]{p['env_key']}=your_key_here[/dim]")
    console.print()


def main():
    args = parse_args()

    console.print()
    console.print(Panel(
        "[bold red]OSINT DEEP BRIEFING ENGINE[/bold red]  "
        "[dim]· Content Writer Edition ·[/dim]\n"
        "[dim]Investigative Analysis · Follow the Money · Expose the Network[/dim]",
        border_style="red", expand=False,
    ))
    console.print()

    if args.providers:
        show_providers()
        return

    # Resolve API key — from .env (loaded at top) or --api-key override
    provider = args.provider
    pinfo = PROVIDERS[provider]
    api_key = args.api_key or os.environ.get(pinfo["env_key"], "")

    if not api_key and args.watchlist != "list":
        console.print(
            f"[bold red]ERROR:[/bold red] No API key for [bold]{provider}[/bold].\n"
            f"  Add to .env:  [cyan]{pinfo['env_key']}=your_key_here[/cyan]\n"
            f"  Get key at:   [cyan]{pinfo['get_key_url']}[/cyan]"
        )
        sys.exit(1)

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Watchlist shortcuts
    if args.watchlist == "list":
        watchlist_list()
        return

    if not args.target:
        console.print("[red]ERROR:[/red] --target is required.")
        sys.exit(1)

    # Collect URLs (optional)
    urls = list(args.urls or [])
    if args.file:
        fp = Path(args.file)
        if not fp.exists():
            console.print(f"[red]ERROR:[/red] File not found: {args.file}")
            sys.exit(1)
        urls += [
            line.strip()
            for line in fp.read_text().splitlines()
            if line.strip() and not line.startswith("#")
        ]

    if args.watchlist == "add":
        if not urls:
            console.print("[red]ERROR:[/red] Provide --urls when using --watchlist add")
            sys.exit(1)
        watchlist_add(args.target, urls)
        return

    if args.watchlist == "run":
        wl = watchlist_load()
        if args.target not in wl:
            console.print(
                f"[red]ERROR:[/red] '{args.target}' not in watchlist. "
                "Run with --watchlist add first."
            )
            sys.exit(1)
        urls = wl[args.target]["urls"]
        console.print(
            f"[dim]Watchlist run for[/dim] [bold]{args.target}[/bold]  "
            f"({len(urls)} URLs)"
        )

    # Dedup URLs
    seen: set[str] = set()
    urls = [u for u in urls if not (u in seen or seen.add(u))]  # type: ignore[func-returns-value]

    console.print(f"[bold]Target:[/bold] {args.target}")
    console.print(
        f"[bold]Provider:[/bold] {pinfo['name']}  |  "
        f"[bold]Mode:[/bold] {args.mode}  |  "
        f"[bold]Depth:[/bold] {args.depth}  |  "
        f"[bold]Angle:[/bold] {args.angle}"
    )
    if urls:
        console.print(f"[bold]URLs provided:[/bold] {len(urls)}")
        for i, u in enumerate(urls):
            console.print(f"  [dim][{i}][/dim] {u}")
    else:
        console.print("[dim]No URLs provided — AI will search the internet automatically[/dim]")
    console.print()

    # ── Phase 1: Scrape (only if URLs provided) ───────────────────────────────
    scraped: list[dict] = []
    if urls:
        console.rule("[dim]PHASE 1 — SCRAPING[/dim]")
        scraped = scrape_all(urls, timeout=args.timeout)
        ok = sum(1 for s in scraped if s["status"] == "ok")
        total_chars = sum(s["char_count"] for s in scraped)
        console.print(f"\n[green]✓[/green] {ok}/{len(urls)} sources · {total_chars:,} chars\n")
        wb_used = [s for s in scraped if s.get("wayback_used")]
        if wb_used:
            console.print(
                f"[dim]ℹ  {len(wb_used)} source(s) fetched via Wayback Machine[/dim]"
            )
        if ok == 0:
            console.print("[bold red]ERROR:[/bold red] All URLs failed to scrape.")
            sys.exit(1)
    else:
        console.print("[dim]Skipping scrape phase — using AI web search[/dim]")

    # ── Phase 2: Build prompt + run briefing ──────────────────────────────────
    console.rule("[dim]PHASE 2 — ANALYSIS[/dim]")
    if scraped:
        prompt = build_briefing_prompt_with_urls(
            args.target, scraped, args.extra, args.angle, args.depth
        )
    else:
        prompt = build_briefing_prompt_no_urls(
            args.target, args.extra, args.angle, args.depth
        )

    briefing, usage = run_briefing(args.target, prompt, provider, api_key)

    # ── Phase 3: Content generation ───────────────────────────────────────────
    console.rule("[dim]PHASE 3 — CONTENT GENERATION[/dim]")
    article_text = thread_text = newsletter_text = legal_text = None

    if args.mode == "article":
        article_text = run_article(briefing, args.target, args.tone, provider, api_key)
    elif args.mode == "thread":
        thread_text = run_thread(briefing, args.target, provider, api_key)
    elif args.mode == "newsletter":
        newsletter_text = run_newsletter(briefing, args.target, provider, api_key)

    if args.legal_review:
        content_for_review = (
            article_text
            or thread_text
            or newsletter_text
            or json.dumps(briefing, indent=2)
        )
        legal_text = run_legal_review(content_for_review, provider, api_key)

    # ── Phase 4: Terminal output ──────────────────────────────────────────────
    console.rule("[dim]PHASE 4 — OUTPUT[/dim]")
    render_terminal(
        args.target, briefing, scraped, usage,
        show_quotes=args.quotes, show_seo=args.seo,
    )

    if article_text:
        console.rule("[bold]ARTICLE DRAFT[/bold]", style="green")
        console.print(article_text)
        console.print()
    if thread_text:
        console.rule("[bold]THREAD[/bold]", style="cyan")
        console.print(thread_text)
        console.print()
    if newsletter_text:
        console.rule("[bold]NEWSLETTER[/bold]", style="magenta")
        console.print(newsletter_text)
        console.print()
    if legal_text:
        console.rule("[bold yellow]LEGAL REVIEW[/bold yellow]", style="yellow")
        console.print(legal_text)
        console.print()

    # ── Watchlist delta ───────────────────────────────────────────────────────
    if args.watchlist == "run":
        wl = watchlist_load()
        runs = wl[args.target].get("runs", [])
        if runs:
            prev_briefing = runs[-1].get("briefing", {})
            delta = compute_delta(prev_briefing, briefing)
            render_delta(args.target, delta)
        runs.append({"at": datetime.now().isoformat(), "briefing": briefing})
        wl[args.target]["runs"] = runs[-10:]
        watchlist_save(wl)

    # ── Save files ────────────────────────────────────────────────────────────
    saved = []
    export = args.export

    if export in ("markdown", "all"):
        p = save_briefing_markdown(
            args.target, briefing, scraped, output_dir,
            include_quotes=args.quotes, include_seo=args.seo,
        )
        saved.append(("Markdown Briefing", p))

    if export in ("json", "all"):
        p = save_json_output(
            args.target, briefing, scraped, output_dir,
            redact_sources=args.redact_sources,
        )
        saved.append(("JSON Data", p))

    if export in ("docx", "all"):
        try:
            p = save_docx(args.target, briefing, scraped, output_dir)
            saved.append(("Word Document", p))
        except ImportError:
            console.print(
                "[yellow]⚠[/yellow]  python-docx not installed — skipping .docx export\n"
                "  Install with: pip install python-docx"
            )

    if article_text:
        saved.append(("Article Draft", save_text(args.target, article_text, "article", output_dir)))
    if thread_text:
        saved.append(("Thread", save_text(args.target, thread_text, "thread", output_dir)))
    if newsletter_text:
        saved.append(("Newsletter", save_text(args.target, newsletter_text, "newsletter", output_dir)))
    if legal_text:
        saved.append(("Legal Review", save_text(args.target, legal_text, "legal_review", output_dir)))

    if saved:
        console.print("[bold]Files saved:[/bold]")
        for label, path in saved:
            console.print(f"  [green]✓[/green] {label}: [cyan]{path.resolve()}[/cyan]")
        console.print()

    console.rule("[dim red]END OF BRIEFING[/dim red]", style="red")
    console.print()


if __name__ == "__main__":
    main()
    if legal_text:
        console.rule("[bold yellow]LEGAL REVIEW[/bold yellow]", style="yellow")
        console.print(legal_text)
        console.print()

    # ── Watchlist delta ───────────────────────────────────────────────────────
    if args.watchlist == "run":
        wl = watchlist_load()
        runs = wl[args.target].get("runs", [])
        if runs:
            prev_briefing = runs[-1].get("briefing", {})
            delta = compute_delta(prev_briefing, briefing)
            render_delta(args.target, delta)
        runs.append({"at": datetime.now().isoformat(), "briefing": briefing})
        wl[args.target]["runs"] = runs[-10:]  # keep last 10 runs
        watchlist_save(wl)

    # ── Save files ────────────────────────────────────────────────────────────
    saved = []
    export = args.export

    if export in ("markdown", "all"):
        p = save_briefing_markdown(args.target, briefing, scraped, output_dir,
                                   include_quotes=args.quotes, include_seo=args.seo)
        saved.append(("Markdown Briefing", p))

    if export in ("json", "all"):
        p = save_json_output(args.target, briefing, scraped, output_dir,
                             redact_sources=args.redact_sources)
        saved.append(("JSON Data", p))

    if export in ("docx", "all"):
        try:
            p = save_docx(args.target, briefing, scraped, output_dir)
            saved.append(("Word Document", p))
        except ImportError:
            console.print("[yellow]⚠[/yellow]  python-docx not installed — skipping .docx export")

    if article_text:
        saved.append(("Article Draft", save_article(args.target, article_text, output_dir)))
    if thread_text:
        saved.append(("Thread", save_thread(args.target, thread_text, output_dir)))
    if newsletter_text:
        saved.append(("Newsletter", save_newsletter(args.target, newsletter_text, output_dir)))
    if legal_text:
        saved.append(("Legal Review", save_legal(args.target, legal_text, output_dir)))

    if saved:
        console.print("[bold]Files saved:[/bold]")
        for label, path in saved:
            console.print(f"  [green]✓[/green] {label}: [cyan]{path.resolve()}[/cyan]")
        console.print()

    console.rule("[dim red]END OF BRIEFING[/dim red]", style="red")
    console.print()


if __name__ == "__main__":
    main()
